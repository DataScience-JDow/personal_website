from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "resume" / "jeremy-dowdy-resume.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(24, 24, 24)
MUTED = RGBColor(85, 85, 85)


def set_run_font(run, size: float, *, bold: bool = False, color: RGBColor = DARK) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_bottom_rule(paragraph, color: str = "D9E2F3") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    list_style = doc.styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_style.font.size = Pt(9.7)
    list_style.paragraph_format.space_after = Pt(2.2)
    list_style.paragraph_format.line_spacing = 1.08
    list_style.paragraph_format.left_indent = Inches(0.22)


def add_header(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("JEREMY DOWDY")
    set_run_font(run, 19, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Data Scientist | Analytics Engineer")
    set_run_font(run, 10.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    contact = (
        "Prosper, TX | jeremydowdy.dev | linkedin.com/in/jeremy-dowdy | "
        "github.com/DataScience-JDow | Jeremydowdy76@gmail.com | 903-474-4110"
    )
    run = p.add_run(contact)
    set_run_font(run, 8.8, color=MUTED)
    add_bottom_rule(p)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    set_run_font(run, 10.6, bold=True, color=BLUE)


def paragraph(doc: Document, text: str, *, size: float = 10.0, after: float = 3) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size)


def label_line(doc: Document, label: str, detail: str, *, size: float = 9.8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(f"{label}: ")
    set_run_font(r, size, bold=True)
    r = p.add_run(detail)
    set_run_font(r, size)


def role(doc: Document, title: str, date_location: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(title)
    set_run_font(r, 10.2, bold=True)
    r = p.add_run(f" | {date_location}")
    set_run_font(r, 9.5, color=MUTED)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    r = p.add_run(text)
    set_run_font(r, 9.65)


def project(doc: Document, title: str, architecture: str, description: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    set_run_font(r, 10.0, bold=True)
    label_line(doc, "Architecture", architecture, size=9.55)
    paragraph(doc, description, size=9.55, after=2)


def build() -> Document:
    doc = Document()
    configure_document(doc)
    add_header(doc)

    heading(doc, "Professional Summary")
    paragraph(
        doc,
        "Highly technical Data Scientist and Analytics Engineer with a Master's in Advanced Data "
        "Analytics and 7+ years of progressive data platform experience at HOLT CAT. Specializes "
        "in enterprise lakehouse architecture, high-throughput ETL/ELT pipelines, predictive "
        "machine learning systems, automated pricing systems, and practical AI-assisted delivery. "
        "Known for translating industrial complexity into trusted data products that improve "
        "pricing, logistics, margin, and executive decision-making.",
        size=10.0,
    )

    heading(doc, "Technical Skills")
    label_line(
        doc,
        "Languages & Frameworks",
        "Python (Pandas, NumPy, Scikit-learn, Seaborn, Matplotlib), SQL (window functions, CTEs), dbt, R, Excel",
    )
    label_line(
        doc,
        "Data Platforms & Warehouses",
        "Microsoft Fabric, Snowflake, Azure, Lakehouse Architecture, Medallion Layering (Bronze/Silver/Gold), PostgreSQL",
    )
    label_line(
        doc,
        "Business Intelligence & Engineering",
        "Power BI (DAX, semantic modeling), Tableau, Alteryx ETL automation, Streamlit",
    )
    label_line(
        doc,
        "DevOps & AI Engineering",
        "Azure DevOps, CI/CD pipelines, Git/GitHub, Cursor, LLM-assisted refactoring, agent workflows",
    )
    label_line(
        doc,
        "Core Competencies",
        "Dimensional modeling, predictive analytics, real-time pricing architecture, stakeholder discovery, API/webhook integration",
    )

    heading(doc, "Professional Experience")
    paragraph(doc, "HOLT CAT 7 yrs 2 mos", size=10.0, after=0)
    role(doc, "Data Scientist (Full-time)", "Sep 2025 - Present | Dallas, TX (Remote)")
    bullet(
        doc,
        "Architected and deployed a machine quoting and market-based pricing application, designing the data pipelines and algorithmic logic behind commercial machinery rate analysis and margin optimization.",
    )
    bullet(
        doc,
        'Spearheading "Project Galaxy," an end-to-end corporate data migration from legacy Snowflake, dbt, and Tableau footprints into a unified Microsoft Fabric Lakehouse environment.',
    )
    bullet(
        doc,
        "Productionized regression, classification, and K-Means clustering systems to predict industrial logistics bottlenecks and optimize fleet distribution networks.",
    )
    bullet(
        doc,
        "Pioneered AI-assisted development workflows using Cursor and LLM agents, accelerating refactoring and pipeline delivery speed by 35%+.",
    )
    bullet(
        doc,
        "Established reusable modeling and evaluation patterns for pricing and operations use cases, improving stakeholder trust in ML-assisted recommendations before production rollout.",
    )

    role(doc, "Data Analyst 3 (Full-time)", "Mar 2023 - Oct 2025 | Irving, TX")
    bullet(
        doc,
        "Engineered scalable Azure and Snowflake ETL/ELT workflows and optimized warehouse indexing to reduce data pipeline ingestion latency by 25%.",
    )
    bullet(
        doc,
        "Conducted warehouse cost auditing, identifying redundant compute clusters and unoptimized queries to reduce cloud architecture overhead.",
    )
    bullet(
        doc,
        "Built production-grade Alteryx workflows to automate operational asset-health metrics, cutting manual reporting by 40%.",
    )
    bullet(
        doc,
        "Partnered directly with senior executive stakeholders to translate ambiguous commercial problems into programmatic technical requirements.",
    )
    bullet(
        doc,
        "Standardized KPI logic and reporting definitions across analytics outputs, reducing interpretation drift in executive and operational reviews.",
    )

    role(doc, "Data Analyst 2", "May 2019 - Mar 2023 | Irving, TX")
    bullet(
        doc,
        "Architected and managed enterprise Power BI and Tableau dashboard suites, providing visibility into billions of historical industrial transaction and telemetry records.",
    )
    bullet(
        doc,
        "Analyzed machinery transaction and telemetry records to isolate operational inefficiency patterns and support annual budget restructuring.",
    )
    bullet(
        doc,
        "Developed dimensional reporting logic and executive-facing views that connected raw operational activity to budget, utilization, and service planning decisions.",
    )

    paragraph(doc, "Fan Cloth 4 yrs 3 mos", size=10.0, after=0)
    role(doc, "Sales Executive", "Nov 2014 - Jan 2019 | Texas")
    bullet(
        doc,
        "Managed regional accounts and revenue targets through pipeline analysis, forecasting metrics, and logistics coordination.",
    )
    bullet(
        doc,
        "Built the business and operational foundation that now informs analytics discovery, product judgment, and stakeholder partnership.",
    )

    heading(doc, "Key Projects")
    project(
        doc,
        "Machine Quoting & Market-Based Pricing Engine",
        "Snowflake, Snowflake Notebooks, Streamlit, XGBoost, Python, SQL.",
        "Architected a pricing intelligence system powered by classification and regression models to replace manual quoting spreadsheets with dynamic machine quoting recommendations for commercial teams, including feature engineering, scenario testing, and analyst-facing review workflows.",
    )
    project(
        doc,
        "Enterprise Lakehouse Migration & Platform Modernization",
        "Microsoft Fabric, Power BI, dbt, Snowflake, SQL, Tableau.",
        "Co-led migration patterns for moving legacy analytics assets into Fabric lakehouse and semantic model layers, standardizing metrics, reusable data contracts, and reporting expectations across business divisions.",
    )
    project(
        doc,
        "Louie Boards Charcuterie CRM & Dashboard",
        "Next.js, Supabase, Vercel, Square API, TypeScript, PostgreSQL.",
        "Built a custom CRM and order workflow for a boutique business, structuring customer history, order details, payment events, and fulfillment status into a durable data product with Square-driven workflow automation and a reporting foundation for repeat-order analysis.",
    )
    project(
        doc,
        "Net Worth & Cashflow Management Engine",
        "Next.js, Supabase, PostgreSQL, Recharts, Vercel Crons.",
        "Built a household finance platform with normalized transaction history, cashflow visibility, scheduled upkeep, caching, and sub-50ms query targets, creating a decision-ready layer for household planning rather than a simple ledger.",
    )

    heading(doc, "Education & Certifications")
    paragraph(doc, "University of North Texas - Master of Science in Advanced Data Analytics, 2018 - 2019", size=9.8, after=1.5)
    paragraph(doc, "East Texas A&M University - Bachelor of Science in Sports Management, 2010 - 2014", size=9.8, after=1.5)
    paragraph(doc, "Microsoft Certified: Fabric Analytics Engineer Associate (DP-600)", size=9.8, after=1.5)
    paragraph(doc, "Active participant in the dbt Global Community; attended Microsoft Fabric Community Conference (FabCon), 2026.", size=9.8, after=1.5)

    return doc


def main() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    props = doc.core_properties
    props.author = "Jeremy Dowdy"
    props.title = "Jeremy Dowdy Resume"
    props.subject = "Resume"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
